"""
doc_handler.py - 文档读取、替换、输出（支持 docx / xlsx / csv / pdf）
"""

import csv
import io
import re
from pathlib import Path

import openpyxl
import pdfplumber
from docx import Document
from openpyxl.styles import Font, PatternFill, Alignment


# ─── 文本提取 ────────────────────────────────────────────────────────────────

def extract_text(file_path: Path) -> str:
    """从文档中提取全部文本，用于发送给 Claude API 识别。"""
    # 提前检查：文件存在且非空
    if not file_path.exists():
        raise ValueError(f"文件不存在：{file_path.name}")
    if file_path.stat().st_size == 0:
        raise ValueError(
            f"上传的文件 [{file_path.name}] 是空文件（0 字节），无法处理。"
            "请确认文件完整后重新上传。"
        )

    suffix = file_path.suffix.lower()
    try:
        if suffix == ".docx":
            return _extract_docx(file_path)
        elif suffix in (".xlsx", ".xls"):
            return _extract_xlsx(file_path)
        elif suffix == ".csv":
            return _extract_csv(file_path)
        elif suffix == ".pdf":
            return _extract_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件格式：{suffix}")
    except ValueError:
        raise
    except Exception as e:
        # 捕获 python-docx 的 PackageNotFoundError 等底层异常，转为可读错误
        err = str(e)
        if "Package not found" in err or "not a zip file" in err.lower():
            raise ValueError(
                f"文件 [{file_path.name}] 不是有效的 Office 文档（可能已损坏、"
                "格式错误，或上传过程中截断）。请重新下载原文件后再上传。"
            )
        raise ValueError(f"读取文件 [{file_path.name}] 时出错：{err}")


def _extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if t:
                        parts.append(t)
    return "\n".join(parts)


def _extract_xlsx(path: Path) -> str:
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for cell in row:
                if cell is not None and str(cell).strip():
                    parts.append(str(cell).strip())
    return "\n".join(parts)


def _extract_csv(path: Path) -> str:
    parts = []
    encodings = ["utf-8-sig", "utf-8", "gbk", "gb2312"]
    for enc in encodings:
        try:
            with open(path, newline="", encoding=enc) as f:
                reader = csv.reader(f)
                for row in reader:
                    for cell in row:
                        if cell.strip():
                            parts.append(cell.strip())
            break
        except (UnicodeDecodeError, Exception):
            continue
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n".join(parts)


# ─── 脱敏应用 ────────────────────────────────────────────────────────────────

def apply_desensitization(input_path: Path, output_path: Path, mapping: dict, original_filename: str):
    """
    将 mapping（{原始文本: 替代内容}）应用到文档，输出脱敏后的文件。
    mapping 中的 key 已按长度降序排列，保证长匹配优先。
    """
    suffix = input_path.suffix.lower()
    if suffix == ".docx":
        _apply_docx(input_path, output_path, mapping)
    elif suffix in (".xlsx", ".xls"):
        _apply_xlsx(input_path, output_path, mapping)
    elif suffix == ".csv":
        _apply_csv(input_path, output_path, mapping)
    elif suffix == ".pdf":
        _apply_pdf_as_docx(input_path, output_path, mapping)
    else:
        raise ValueError(f"不支持的文件格式：{suffix}")


def _sorted_mapping(mapping: dict) -> list:
    """按 key 长度降序，避免短字符串先替换导致长字符串无法匹配。"""
    return sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True)


def _replace_text(text: str, sorted_pairs: list) -> str:
    """
    两阶段替换，防止级联污染：
    第一阶段：将每个 original 替换为唯一占位符（\x00DEIDENT_N\x00）；
    第二阶段：将所有占位符替换为最终 replacement。
    这样 entity A 的替换结果不会被 entity B 再次命中。
    """
    placeholder_map: dict = {}  # placeholder → replacement
    for i, (orig, repl) in enumerate(sorted_pairs):
        if orig in text:
            marker = f"\x00DEIDENT_{i}\x00"
            text = text.replace(orig, marker)
            placeholder_map[marker] = repl
    for marker, repl in placeholder_map.items():
        text = text.replace(marker, repl)
    return text


def tokenize_text_with_mapping(text: str, mapping: dict) -> str:
    """将原文按映射表替换为脱敏 token/替代值，适用于发送给外部 AI。"""
    if not text or not mapping:
        return text
    return _replace_text(text, _sorted_mapping(mapping))


def restore_text_with_mapping(text: str, reverse_mapping: dict) -> str:
    """将外部 AI 返回文本按反向映射还原为原文。"""
    if not text or not reverse_mapping:
        return text
    return _replace_text(text, _sorted_mapping(reverse_mapping))


def _apply_docx(input_path: Path, output_path: Path, mapping: dict):
    doc = Document(str(input_path))
    pairs = _sorted_mapping(mapping)

    def process_paragraph(para):
        # 先尝试在每个 run 内替换
        full_text_before = "".join(r.text for r in para.runs)
        for run in para.runs:
            if run.text:
                run.text = _replace_text(run.text, pairs)
        full_text_after = "".join(r.text for r in para.runs)

        # 若替换后合并文本与预期不符（说明有实体跨 run），则合并重写
        expected = _replace_text(full_text_before, pairs)
        if full_text_after != expected:
            # 合并所有 run 到第一个，保留其字体格式
            if para.runs:
                first_run = para.runs[0]
                first_run.text = expected
                for run in para.runs[1:]:
                    run.text = ""

    for para in doc.paragraphs:
        process_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)

    doc.save(str(output_path))


def _apply_xlsx(input_path: Path, output_path: Path, mapping: dict):
    wb = openpyxl.load_workbook(str(input_path))
    pairs = _sorted_mapping(mapping)

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    cell.value = _replace_text(cell.value, pairs)

    wb.save(str(output_path))


def _apply_csv(input_path: Path, output_path: Path, mapping: dict):
    pairs = _sorted_mapping(mapping)
    encodings = ["utf-8-sig", "utf-8", "gbk"]
    rows = []
    used_enc = "utf-8"

    for enc in encodings:
        try:
            with open(input_path, newline="", encoding=enc) as f:
                rows = list(csv.reader(f))
            used_enc = enc
            break
        except UnicodeDecodeError:
            continue

    new_rows = []
    for row in rows:
        new_rows.append([_replace_text(cell, pairs) for cell in row])

    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerows(new_rows)


def _apply_pdf_as_docx(input_path: Path, output_path: Path, mapping: dict):
    """PDF 提取文本后替换，输出为 DOCX（格式简化）。"""
    pairs = _sorted_mapping(mapping)
    doc = Document()

    # 添加说明段落
    note = doc.add_paragraph()
    note_run = note.add_run("【说明】本文档由 PDF 格式转换而来，原始排版可能有所简化。")
    note_run.bold = True
    doc.add_paragraph("")  # 空行

    with pdfplumber.open(str(input_path)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            replaced = _replace_text(text, pairs)
            if i > 0:
                doc.add_page_break()
            for line in replaced.split("\n"):
                doc.add_paragraph(line)

    # 调整输出路径后缀为 .docx
    if output_path.suffix.lower() == ".pdf":
        output_path = output_path.with_suffix(".docx")

    doc.save(str(output_path))
    return output_path


# ─── 脱敏还原 ────────────────────────────────────────────────────────────────

def restore_document(input_path: Path, output_path: Path, reverse_mapping: dict):
    """使用反向映射表将脱敏文档还原。"""
    apply_desensitization(input_path, output_path, reverse_mapping, input_path.name)


# ─── 映射表 Excel ────────────────────────────────────────────────────────────

def create_mapping_excel(entities: list, output_path: Path):
    """生成映射表 Excel（类别 / 原始内容 / 替代内容）。"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "脱敏映射表"

    # 表头样式
    header_fill = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    headers = ["类别", "原始内容", "替代内容"]

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # 数据行
    row_fill_even = PatternFill(start_color="EFF6FF", end_color="EFF6FF", fill_type="solid")
    for i, entity in enumerate(entities, 2):
        ws.cell(row=i, column=1, value=entity.get("category", ""))
        ws.cell(row=i, column=2, value=entity.get("original", ""))
        ws.cell(row=i, column=3, value=entity.get("replacement", ""))
        if i % 2 == 0:
            for col in range(1, 4):
                ws.cell(row=i, column=col).fill = row_fill_even

    # 列宽
    ws.column_dimensions["A"].width = 20
    ws.column_dimensions["B"].width = 40
    ws.column_dimensions["C"].width = 40

    # 冻结首行
    ws.freeze_panes = "A2"

    wb.save(str(output_path))


def read_mapping_excel(mapping_path: Path) -> dict:
    """
    读取映射表 Excel，返回 {替代内容: 原始内容} 的反向映射（用于还原）。
    同时也返回正向映射，由调用方选择。
    """
    wb = openpyxl.load_workbook(str(mapping_path), read_only=True, data_only=True)
    ws = wb.active

    reverse_mapping = {}
    rows = list(ws.iter_rows(values_only=True))

    if not rows:
        return {}

    # 找到列索引（兼容表头顺序变化）
    header = [str(h).strip() if h else "" for h in rows[0]]
    try:
        orig_col = header.index("原始内容")
        repl_col = header.index("替代内容")
    except ValueError:
        # 默认：第2列=原始, 第3列=替代
        orig_col, repl_col = 1, 2

    for row in rows[1:]:
        if len(row) <= max(orig_col, repl_col):
            continue
        original = str(row[orig_col]).strip() if row[orig_col] else ""
        replacement = str(row[repl_col]).strip() if row[repl_col] else ""
        if original and replacement:
            reverse_mapping[replacement] = original

    return reverse_mapping


def get_output_suffix(original_suffix: str) -> str:
    """PDF 输入输出为 docx，其他保持原格式。"""
    if original_suffix.lower() == ".pdf":
        return ".docx"
    return original_suffix.lower()
